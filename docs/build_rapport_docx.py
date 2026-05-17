#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Génère le rapport PFA Tariki au format Microsoft Word (.docx).
Usage : python3 build_rapport_docx.py
"""

from __future__ import annotations

import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "rapport_assets"
OUT_DOCX = ROOT / "RAPPORT-TARIKI-PFA.docx"


# ─── Utilitaires Word ─────────────────────────────────────────────────────────

def set_cell_shading(cell, fill: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_field_toc(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    paragraph.add_run(
        "Cliquez avec le bouton droit sur cette zone → « Mettre à jour les champs » "
        "pour afficher la table des matières."
    ).italic = True


def setup_document(doc: Document):
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)
    for level, size in [(1, 16), (2, 14), (3, 12)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)


def add_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Tariki — Rapport PFA 2026  |  Page ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def p(doc, text):
    doc.add_paragraph(text)


def pb(doc):
    doc.add_page_break()


def fig(doc, name: str, caption: str, width_cm=15.5):
    path = ASSETS / name
    if not path.exists():
        p(doc, f"[Figure manquante : {name}]")
        return
    doc.add_picture(str(path), width=Cm(width_cm))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(10)
    doc.add_paragraph()


def table(doc, headers, rows, caption: str | None = None):
    if caption:
        c = doc.add_paragraph(caption)
        c.runs[0].bold = True
        c.runs[0].font.size = Pt(10)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], "1E40AF")
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for run in cells[ci].paragraphs[0].runs:
                run.font.size = Pt(10)
        if ri % 2 == 1:
            for c in cells:
                set_cell_shading(c, "F8FAFC")
    doc.add_paragraph()


def cover_page(doc: Document):
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph("TARIKI")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.size = Pt(36)
    t.runs[0].font.bold = True
    t.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    sub = doc.add_paragraph(
        "Système intelligent de gestion du trafic urbain\nGrand Casablanca — Royaume du Maroc"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    doc.add_paragraph()
    doc.add_paragraph()
    lines = [
        "Projet de fin d'années (PFA)",
        "",
        "Réalisé par : Kalil",
        "Compte GitHub : kalil-cyber",
        "",
        "Encadrant : …………………………………………",
        "Établissement / Filière : …………………………………………",
        "",
        "Année universitaire 2025–2026",
        "",
        "Dépôt source : https://github.com/kalil-cyber/PFA-4eme",
    ]
    for line in lines:
        para = doc.add_paragraph(line)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if line and para.runs:
            para.runs[0].font.size = Pt(12)
    pb(doc)


def front_matter(doc: Document):
    h(doc, "Dédicace", 1)
    p(
        doc,
        "À ma famille, à mes proches, et à tous les usagers de la rocade casaouie "
        "qui méritent des routes plus fluides — ce travail leur est dédié.",
    )
    pb(doc)

    h(doc, "Remerciements", 1)
    p(
        doc,
        "Je remercie mon encadrant pédagogique pour la disponibilité et les retours "
        "constructifs, l'équipe enseignante de la filière, les camarades ayant testé "
        "l'application, ainsi que la communauté open source (React, Node.js, OpenStreetMap) "
        "sans laquelle un projet de cette envergure en solo serait impossible.",
    )
    pb(doc)

    h(doc, "Résumé", 1)
    p(
        doc,
        "Tariki est une plateforme web de gestion intelligente de la circulation pour "
        "Casablanca. Elle intègre un dataset de mobilité dérivé de trajectoires type Waze "
        "(environ 440 observations par jour, sept jours de la semaine), une API temps réel "
        "(Node.js, Express, Socket.io), des interfaces conducteur et administrateur, des "
        "prévisions par régression linéaire, un assistant conversationnel, et un module "
        "Webcams & surveillance alimenté par un référentiel local de 25 points géolocalisés. "
        "Le système est déployable sur GitHub, Render et Vercel, et fonctionne en mode "
        "démonstration sans base PostgreSQL grâce à un store mémoire.",
    )
    p(
        doc,
        "Mots-clés : ville intelligente ; ITS ; trafic urbain ; prédiction ; crowdsourcing ; "
        "React ; Node.js ; Casablanca ; mobilité durable.",
    )
    pb(doc)

    h(doc, "Abstract", 1)
    p(
        doc,
        "Tariki is a web-based intelligent traffic management platform for Casablanca, Morocco. "
        "It combines cleaned crowdsourced mobility data, a real-time REST/WebSocket API, "
        "role-based user interfaces, short-term congestion forecasting, a contextual chatbot, "
        "and a local surveillance dataset. The architecture supports both production "
        "(PostgreSQL) and demo (in-memory) modes. This report documents context, scientific "
        "background, design, implementation, tests, and perspectives.",
    )
    p(doc, "Keywords: smart city, ITS, traffic prediction, floating car data, full-stack web.")
    pb(doc)

    h(doc, "Table des matières", 1)
    add_field_toc(doc.add_paragraph())
    pb(doc)

    h(doc, "Liste des figures", 1)
    figures = [
        "Architecture générale en couches",
        "Flux de données temps réel",
        "Diagramme de cas d'utilisation",
        "Séquence d'authentification",
        "Courbe de prédiction (régression linéaire)",
        "Segments routiers sur Casablanca",
        "Architecture de déploiement",
        "Planning projet (Gantt)",
        "Modèle de données simplifié",
        "Maquettes des interfaces",
    ]
    for i, leg in enumerate(figures, 1):
        p(doc, f"Figure {i} — {leg}")
    pb(doc)

    h(doc, "Liste des tableaux", 1)
    tables = [
        "Indicateurs urbains de Casablanca",
        "Composants ITS vs équivalents Tariki",
        "Comparaison des méthodes de prédiction",
        "Stack technique du projet",
        "Fichiers du dataset Waze",
        "Catégories du dataset surveillance",
        "Endpoints API principaux",
        "Objectifs spécifiques et critères",
        "Comptes de démonstration",
        "Risques et mitigations",
    ]
    for i, leg in enumerate(tables, 1):
        p(doc, f"Tableau {i} — {leg}")
    pb(doc)


def chapter_intro(doc):
    h(doc, "Introduction générale", 1)
    p(
        doc,
        "La métropole de Casablanca concentre une part majeure de l'activité économique "
        "marocaine. Chaque jour, des centaines de milliers de déplacements génèrent des "
        "congestions aux heures de pointe, un surcoût en carburant, du stress et des émissions "
        "de CO₂. Les collectivités et les opérateurs de transport cherchent des outils numériques "
        "capables d'observer, d'anticiper et de communiquer sur l'état du réseau routier.",
    )
    p(
        doc,
        "Les systèmes de transport intelligents (ITS) historiques reposaient sur des capteurs "
        "coûteux : boucles magnétiques, caméras fixes, stations de pesage. Aujourd'hui, les "
        "données participatives issues d'applications mobiles (Waze, Google Maps) et l'open data "
        "cartographique (OpenStreetMap) ouvrent une voie complémentaire à moindre coût, "
        "particulièrement pertinente pour un projet étudiant de fin d'années (PFA).",
    )
    p(
        doc,
        "Tariki — nom du projet — est une plateforme full-stack qui démontre cette intégration : "
        "chargement d'un dataset Waze nettoyé, simulation temps réel sur une dizaine de segments "
        "nommés, API REST et WebSocket, portails conducteur et administrateur, module de prédiction, "
        "gestion d'incidents, chatbot contextuel et page Webcams & surveillance. Ce rapport en "
        "documente le contexte, l'état de l'art, la conception, la réalisation, les tests et les "
        "perspectives, en s'appuyant sur le dépôt GitHub public kalil-cyber/PFA-4eme.",
    )
    fig(doc, "fig01_architecture_generale.png", "Figure 1 — Architecture générale du système Tariki")
    p(
        doc,
        "Structure du document : la partie I pose le contexte casablancais et la problématique ; "
        "la partie II présente une revue scientifique (smart cities, ITS, crowdsourcing, prédiction) ; "
        "la partie III détaille l'analyse et la conception ; la partie IV traite des données ; "
        "la partie V la réalisation technique ; la partie VI le bilan, les limites et la conclusion.",
    )
    pb(doc)


def chapter_contexte(doc):
    h(doc, "Partie I — Contexte et problématique", 1)
    h(doc, "Casablanca : enjeux de mobilité", 2)
    p(
        doc,
        "Casablanca est la plus grande ville du Maroc et un pôle portuaire, industriel et "
        "tertiaire. Son réseau routier combine des axes structurants (boulevard Mohammed V, "
        "Zerktouni, corniche Aïn Diab, rocade), des accès au port, et des liaisons avec les "
        "autoroutes à péage gérées par les Autoroutes du Maroc (ADM). Le tramway (lignes T1, T2) "
        "et le réseau de bus Casa Transport offrent une alternative, mais la congestion aux "
        "heures de pointe limite l'intermodalité.",
    )
    table(
        doc,
        ["Indicateur", "Ordre de grandeur", "Impact sur le trafic"],
        [
            ["Population agglomération", "Plusieurs millions d'habitants", "Forte demande de déplacement"],
            ["Motorisation", "Dominante individuelle", "Saturation des axes centraux"],
            ["Tramway T1 / T2", "Lignes structurantes", "Report modal partiel"],
            ["Smartphone", "Forte pénétration", "Crowdsourcing viable (Waze)"],
            ["Port / logistique", "Flux poids lourds", "Goulets périphériques"],
        ],
        "Tableau 1 — Indicateurs urbains (ordre de grandeur)",
    )
    h(doc, "Problématique scientifique et opérationnelle", 2)
    p(
        doc,
        "La question centrale de ce PFA est la suivante : peut-on concevoir, avec des moyens "
        "limités d'étudiant mais une stack moderne, une plateforme crédible d'information trafic "
        "pour Casablanca ? Cette question se décline en sous-problèmes : visualisation quasi temps "
        "réel, gestion d'incidents, prédiction à court terme, centralisation de sources hétérogènes "
        "(dataset, météo, POI, surveillance), et séparation des rôles conducteur / administrateur.",
    )
    table(
        doc,
        ["Question", "Enjeu", "Réponse Tariki"],
        [
            ["État du trafic ?", "Décision conducteur", "Carte + WebSocket 3 s"],
            ["Incidents ?", "Sécurité", "CRUD + notifications"],
            ["Anticipation ?", "Planification", "Régression linéaire 30 min"],
            ["Données hétérogènes ?", "Smart city", "API discover + datasets locaux"],
        ],
        "Tableau 2 — Problématique déclinée",
    )
    h(doc, "Hypothèses de travail", 2)
    p(doc, "H1 — Un dataset Waze nettoyé (bbox Casablanca, 7 jours) suffit à alimenter une démonstration réaliste.")
    p(doc, "H2 — Une régression linéaire sur fenêtre glissante est suffisante pour un MVP pédagogique de prédiction.")
    p(doc, "H3 — Un dataset local de surveillance (25 points) améliore l'UX sans dépendre d'API externes instables.")
    p(doc, "H4 — L'architecture React / Node / Socket.io est adaptée au déploiement cloud gratuit (Render, Vercel).")
    pb(doc)


def chapter_etat_art(doc):
    h(doc, "Partie II — État de l'art et revue scientifique", 1)
    h(doc, "Villes intelligentes (Smart Cities)", 2)
    p(
        doc,
        "Batty (2013) et Bibri & Krogstie (2017) décrivent l'intégration des infrastructures "
        "physiques et numériques pour améliorer qualité de vie, efficacité énergétique et mobilité. "
        "Les piliers pertinents pour Tariki sont : capteurs et IoT (feux, caméras), plateformes de "
        "données (APIs ouvertes), et applications citoyennes (information, participation). Tariki "
        "se positionne comme brique logicielle de démonstration, non comme infrastructure municipale.",
    )
    h(doc, "Systèmes de transport intelligents (ITS)", 2)
    table(
        doc,
        ["Composant ITS", "Exemple industriel", "Équivalent Tariki"],
        [
            ["Acquisition", "Détecteurs, GPS flotte", "Dataset Waze + simulateur"],
            ["Stockage", "Data lake, SCADA", "PostgreSQL / memoryStore"],
            ["Analyse", "Modèles prédictifs", "predictionEngine.js"],
            ["Diffusion", "PMV, apps", "React + Socket.io"],
            ["Contrôle", "Feux adaptatifs", "Non implémenté (perspective)"],
        ],
        "Tableau 3 — Composants ITS",
    )
    fig(doc, "fig02_flux_donnees.png", "Figure 2 — Flux de données temps réel")
    h(doc, "Crowdsourcing et Floating Car Data", 2)
    p(
        doc,
        "Goodchild (2007) introduit le concept de « citoyens capteurs » : les usagers "
        "deviennent producteurs de données géolocalisées. Waze et Google Maps illustrent le "
        "Floating Car Data (FCD). Avantages : couverture spatiale large, mise à jour fréquente, "
        "faible coût d'infrastructure. Limites : biais d'échantillonnage (conducteurs smartphone), "
        "qualité variable, enjeux de vie privée. Le dataset tariki_cleaned_dataset du projet "
        "repose sur ce paradigme, avec tables par jour (lundi à dimanche).",
    )
    h(doc, "Prédiction du trafic", 2)
    p(
        doc,
        "Vlahogianni et al. (2014) recensent les défis de la prédiction à court terme (5–60 min). "
        "Les approches vont des modèles statistiques (ARIMA) au deep learning (LSTM, GNN). "
        "Lv et al. (2015) montrent que les LSTM capturent mieux les non-linéarités. Tariki "
        "implémente en MVP une régression linéaire sur les 24 derniers points de congestion, "
        "avec coefficient R² renvoyé au client pour indiquer la confiance.",
    )
    table(
        doc,
        ["Méthode", "Forces", "Faiblesses", "Usage Tariki"],
        [
            ["Régression linéaire", "Simple, interprétable", "Non-linéarité", "MVP actuel"],
            ["ARIMA", "Séries temporelles", "Stationnarité", "Perspective"],
            ["LSTM", "Mémoire longue", "Données, GPU", "backend/python/"],
            ["GNN", "Réseau routier", "Complexité", "Recherche future"],
        ],
        "Tableau 4 — Méthodes de prédiction",
    )
    fig(doc, "fig05_prediction.png", "Figure 5 — Exemple de courbe de prédiction")
    fig(doc, "fig03_cas_utilisation.png", "Figure 3 — Diagramme de cas d'utilisation")
    h(doc, "Vidéosurveillance et vision", 2)
    p(
        doc,
        "Les travaux sur YOLO et OpenCV permettent d'estimer densité et files d'attente à partir "
        "de caméras fixes. Les péages ADM exposent parfois des flux publics. Tariki intègre "
        "surveillance_casablanca.json (9 webcams, 4 péages, 8 feux, 4 zones) pour la "
        "géolocalisation et l'UX, sans traitement vidéo automatique en Phase 1.",
    )
    pb(doc)


def chapter_conception(doc):
    h(doc, "Partie III — Analyse et conception", 1)
    h(doc, "Objectifs et critères de réussite", 2)
    table(
        doc,
        ["ID", "Objectif", "Critère"],
        [
            ["O1", "Visualiser le trafic", "Segments colorés, WS 3 s"],
            ["O2", "Gérer incidents", "CRUD + broadcast"],
            ["O3", "Prédire congestion", "GET /api/predictions"],
            ["O4", "Rôles admin / user", "JWT + routes protégées"],
            ["O5", "Dataset multi-jours", "7 JSON + apply-day"],
            ["O6", "Assistant", "Chatbot TrafficChatbot"],
            ["O7", "Documenter / déployer", "README, GitHub, Render"],
        ],
        "Tableau 5 — Objectifs spécifiques",
    )
    h(doc, "Méthodologie", 2)
    p(doc, "1. Analyse du besoin (conducteur vs opérateur). 2. Préparation des données (Excel → JSON). 3. Conception UML simplifiée. 4. Implémentation par couches (API puis UI). 5. Tests smoke (test:api). 6. Déploiement et rapport.")
    fig(doc, "fig04_sequence_connexion.png", "Figure 4 — Diagramme de séquence : authentification")
    fig(doc, "fig09_modele_donnees.png", "Figure 9 — Modèle de données simplifié")
    h(doc, "Architecture logicielle", 2)
    table(
        doc,
        ["Couche", "Technologie", "Rôle"],
        [
            ["Présentation", "React 18, Vite, Tailwind", "UI, cartes, dashboards"],
            ["API", "Express, JWT, bcrypt", "REST sécurisé"],
            ["Temps réel", "Socket.io", "traffic:update, incidents"],
            ["Métier", "Simulateur, prédiction, chat", "Logique applicative"],
            ["Données", "PostgreSQL / mémoire", "Persistance"],
            ["Carto", "Mapbox GL, Leaflet", "Visualisation"],
        ],
        "Tableau 6 — Stack technique",
    )
    p(
        doc,
        "Structure des répertoires : datasets/ (Waze + surveillance), backend/src/ (routes, services, db), "
        "frontend/src/ (pages, components, contexts), docs/ (documentation et ce rapport), scripts/ (publication GitHub).",
    )
    fig(doc, "fig10_interfaces.png", "Figure 10 — Maquettes des interfaces principales")
    pb(doc)


def chapter_donnees(doc):
    h(doc, "Partie IV — Données", 1)
    h(doc, "Dataset trafic Waze (tariki_cleaned_dataset)", 2)
    p(
        doc,
        "Le cœur métier repose sur sept fichiers JSON (table_5_monday à table_11_sunday), "
        "des tables annexes (population, transports, types de routes, usage du sol) et un résumé CSV. "
        "Le service tarikiDatasetLoader.js extrait les coordonnées dans la bounding box de Casablanca, "
        "génère environ 13 segments nommés (Mohammed V, Zerktouni, Corniche, etc.) et dérive un "
        "niveau de congestion 0–100 % par seed déterministe et index du jour.",
    )
    table(
        doc,
        ["Fichier / ressource", "Contenu"],
        [
            ["table_5_monday.json … sunday", "Coordonnées et congestion par jour"],
            ["table_0_coordinates", "Référentiel spatial"],
            ["table_1 … table_4", "Données socio-transport"],
            ["dataset_summary.csv", "Synthèse statistique"],
        ],
        "Tableau 7 — Fichiers dataset Waze",
    )
    fig(doc, "fig06_carte_segments.png", "Figure 6 — Répartition schématique des segments")
    h(doc, "Dataset surveillance (surveillance_casablanca.json)", 2)
    table(
        doc,
        ["Catégorie", "Nombre", "Rôle"],
        [
            ["Webcam", "9", "Points de vue référencés"],
            ["Péage", "4", "Passages autoroutiers ADM"],
            ["Feu", "8", "Carrefours signalés"],
            ["Surveillance", "4", "Zones fixes (tunnel, marina…)"],
        ],
        "Tableau 8 — Dataset surveillance (25 points)",
    )
    p(
        doc,
        "Chargement 100 % local : la page /webcams s'affiche sans latence liée à des APIs "
        "externes (ADM, MapCam, Overpass). Le frontend embarque une copie via utils/surveillance.js.",
    )
    h(doc, "Sources externes optionnelles", 2)
    p(doc, "Open-Meteo (météo), OSM/Overpass (POI, dégradé), OpenAI (chat enrichi si clé API). Mode dégradé garanti sans clés tierces.")
    pb(doc)


def chapter_realisation(doc):
    h(doc, "Partie V — Réalisation technique", 1)
    h(doc, "Backend Node.js", 2)
    p(
        doc,
        "Le serveur Express (port 4000) expose les routes : /api/auth (login, register, check-email, me), "
        "/api/traffic, /api/incidents, /api/predictions, /api/dataset, /api/discover, /api/chat, /api/logs, /api/alerts. "
        "Le simulateur trafficSimulator.js met à jour la congestion toutes les 3 secondes. "
        "Socket.io diffuse traffic:update et prediction:update.",
    )
    table(
        doc,
        ["Méthode", "Route", "Description"],
        [
            ["GET", "/api/health", "État du service"],
            ["POST", "/api/auth/login", "Connexion JWT"],
            ["GET", "/api/auth/check-email", "Reconnaissance compte"],
            ["GET", "/api/traffic/roads", "Segments routiers"],
            ["GET", "/api/predictions", "Prévisions 30 min"],
            ["CRUD", "/api/incidents", "Gestion incidents"],
            ["GET", "/api/discover/webcams", "Points surveillance"],
            ["POST", "/api/chat/message", "Assistant"],
        ],
        "Tableau 9 — Endpoints API principaux",
    )
    h(doc, "Frontend React", 2)
    p(
        doc,
        "Pages : Accueil, Conducteur (/driver), Carte (/carte), Connexion (/connexion), "
        "Webcams, Dashboard admin (/tariki-ops), Aide. Composants : cartes Mapbox/Leaflet, "
        "graphiques Recharts, TrafficChatbot flottant, PredictionAnalystPanel. AuthPage : "
        "toast + connexion automatique si email déjà enregistré (check-email).",
    )
    h(doc, "Intelligence artificielle", 2)
    p(
        doc,
        "predictionEngine.js : moindres carrés sur 24 points, horizon 6×5 min. predictionAnalyst.js "
        "synthétise les zones à risque. Extension LSTM documentée dans backend/python/ (FastAPI prévu).",
    )
    h(doc, "Sécurité", 2)
    p(doc, "JWT signé, rôles admin/user, bcrypt, code admin (ADMIN_ACCESS_CODE), CORS, routes protégées par authMiddleware et adminMiddleware.")
    h(doc, "Déploiement", 2)
    fig(doc, "fig07_deploiement.png", "Figure 7 — Architecture de déploiement (GitHub, Vercel, Render)")
    fig(doc, "fig08_gantt.png", "Figure 8 — Planning projet (diagramme de Gantt)")
    p(doc, "Commandes : npm run install:all ; npm run dev:backend ; npm run dev:frontend. Variables : JWT_SECRET, VITE_API_URL, DATABASE_URL, USE_MEMORY=true pour démo.")
    h(doc, "Tests et validation", 2)
    p(doc, "Smoke test backend : npm run test:api. Tests manuels : connexion, carte temps réel, CRUD incidents, sélecteur jour dataset, page webcams, dashboard KPIs.")
    pb(doc)


def chapter_bilan(doc):
    h(doc, "Partie VI — Bilan, limites et conclusion", 1)
    h(doc, "Résultats obtenus", 2)
    p(
        doc,
        "Le projet a livré une plateforme fonctionnelle en local et déployable sur le cloud, "
        "un dépôt GitHub versionné, une documentation README complète, un dataset documenté, "
        "et ce rapport. Les objectifs O1 à O7 sont atteints au niveau MVP.",
    )
    h(doc, "Limites", 2)
    p(doc, "La simulation ne remplace pas des capteurs réels. La prédiction linéaire est simpliste. Les webcams n'embarquent pas toutes un flux vidéo. Les feux sont des positions de référence sans état dynamique municipal.")
    table(
        doc,
        ["Risque", "Mitigation"],
        [
            ["Qualité dataset", "BBox Casa, documentation traçabilité"],
            ["Sécurité JWT", "Secrets, HTTPS en production"],
            ["APIs tierces", "Datasets locaux, mode dégradé"],
            ["Cold start Render", "Health check, patience utilisateur"],
        ],
        "Tableau 10 — Risques et mitigations",
    )
    h(doc, "Perspectives", 2)
    p(doc, "1. Intégration LSTM / Prophet. 2. Partenariat API trafic officielle. 3. Application mobile React Native. 4. GTFS tramway Casa. 5. Vision YOLO sur flux péage. 6. Feux connectés si API municipale disponible.")
    h(doc, "Conclusion générale", 2)
    p(
        doc,
        "Tariki démontre la faisabilité d'une smart-mobility demo crédible pour Casablanca, "
        "en s'inscrivant dans les champs des ITS, des smart cities et du machine learning appliqué "
        "au transport. Le jury peut évaluer le code source public, lancer l'application en local, "
        "et reproduire les scénarios décrits en annexes. Ce travail ouvre une trajectoire claire "
        "vers des modèles deep learning et la vidéosurveillance intelligente.",
    )
    pb(doc)


def chapter_biblio_annexes(doc):
    h(doc, "Bibliographie", 1)
    refs = [
        "Batty, M. (2013). The New Science of Cities. MIT Press.",
        "Bibri, S. E., & Krogstie, J. (2017). Smart sustainable cities of the future. Energy Procedia.",
        "Dimitriou, H. T., & Gakenheimer, M. (2011). Urban Transport in the Developing World. Edward Elgar.",
        "Goodchild, M. F. (2007). Citizens as sensors. GeoJournal.",
        "Lv, Y., et al. (2015). Traffic flow prediction with big data. IEEE TITS.",
        "Vlahogianni, E. I., et al. (2014). Short-term traffic forecasting. Transportation Research Part C.",
        "Bar-Gera, H. (2017). Traffic assignment. Transportation Research Part B.",
        "Documentation React, Express, Socket.io, PostgreSQL, Mapbox GL.",
        "OpenStreetMap contributors ; ADM Autoroutes du Maroc.",
    ]
    for i, r in enumerate(refs, 1):
        p(doc, f"[{i}] {r}")
    pb(doc)

    h(doc, "Annexes", 1)
    h(doc, "Annexe A — Comptes de démonstration", 2)
    table(
        doc,
        ["Rôle", "Email", "Mot de passe"],
        [
            ["Administrateur", "kalil@gmail.com", "0000"],
            ["Conducteur", "kpl@gmail.com", "0000"],
            ["Code admin", "—", "0000 (env)"],
        ],
    )
    h(doc, "Annexe B — URLs locales", 2)
    table(
        doc,
        ["Service", "URL"],
        [
            ["Application", "http://localhost:5173"],
            ["Connexion", "http://localhost:5173/connexion"],
            ["API health", "http://localhost:4000/api/health"],
            ["Webcams", "http://localhost:5173/webcams"],
            ["Admin", "http://localhost:5173/tariki-ops"],
        ],
    )
    h(doc, "Annexe C — Fichiers clés du dépôt", 2)
    table(
        doc,
        ["Fichier", "Rôle"],
        [
            ["tarikiDatasetLoader.js", "Chargement dataset Waze"],
            ["predictionEngine.js", "Régression linéaire"],
            ["trafficSimulator.js", "Temps réel 3 s"],
            ["surveillance_casablanca.json", "25 points surveillance"],
            ["AuthPage.jsx", "Authentification unifiée"],
            ["WebcamsPage.jsx", "Carte Leaflet surveillance"],
        ],
    )
    h(doc, "Annexe D — Variables d'environnement", 2)
    p(doc, "JWT_SECRET, ADMIN_ACCESS_CODE, DATABASE_URL, USE_MEMORY, CORS_ORIGIN, VITE_API_URL, VITE_WS_URL, OPENAI_API_KEY (optionnel).")
    h(doc, "Annexe E — Événements WebSocket", 2)
    p(doc, "traffic:update, prediction:update, incident:created, incident:updated, incident:deleted, simulation:toggle.")
    h(doc, "Annexe F — Glossaire", 2)
    table(
        doc,
        ["Terme", "Définition"],
        [
            ["ITS", "Intelligent Transportation System"],
            ["FCD", "Floating Car Data"],
            ["JWT", "JSON Web Token"],
            ["MVP", "Minimum Viable Product"],
            ["LSTM", "Long Short-Term Memory"],
            ["POI", "Point d'intérêt"],
            ["ADM", "Autoroutes du Maroc"],
        ],
    )
    h(doc, "Annexe G — Guide captures d'écran pour le jury", 2)
    p(
        doc,
        "Pour compléter le rapport remis au jury, insérer ici des captures d'écran : (1) page d'accueil, "
        "(2) carte conducteur avec segments colorés, (3) dashboard administrateur avec KPIs, "
        "(4) page connexion avec notification auto-login, (5) module webcams avec filtres, "
        "(6) panneau prédiction analyste. Dans Word : Insertion → Images → sélectionner les PNG.",
    )


def chapter_detail_modules(doc):
    """Sections détaillées par module (volume rapport PFA)."""
    h(doc, "Partie V bis — Détail des modules implémentés", 1)

    modules = [
        (
            "Module authentification (auth.js, AuthPage.jsx)",
            "L'authentification unifiée sur /connexion gère l'inscription et la connexion des conducteurs "
            "et des administrateurs. Le endpoint GET /api/auth/check-email permet, dès la saisie de l'adresse, "
            "de détecter un compte existant : le frontend affiche une notification et peut déclencher une "
            "connexion automatique après saisie du mot de passe. Les mots de passe sont hashés avec bcrypt ; "
            "le JWT porte le rôle (admin ou user) et une expiration configurable. Le code administrateur "
            "(ADMIN_ACCESS_CODE) protège la création de comptes à privilèges. Ce choix répond au besoin UX "
            "identifié en tests utilisateurs tout en conservant un modèle stateless adapté au déploiement Render.",
        ),
        (
            "Module trafic et simulateur (trafficSimulator.js, traffic.js)",
            "En l'absence de flux capteurs municipaux en temps réel, le simulateur actualise environ treize "
            "segments toutes les trois secondes avec des niveaux de congestion entre 0 et 100 %. Les variations "
            "sont pseudo-aléatoires mais bornées pour éviter des sauts irréalistes. Chaque mise à jour est "
            "émise via Socket.io sous l'événement traffic:update ; les hooks React useTrafficData et "
            "useTrafficNotifications répercutent les changements sur la carte et les toasts. L'administrateur "
            "peut activer ou désactiver la simulation via POST /api/traffic/simulation. Lors de l'application "
            "d'un jour du dataset (POST /api/dataset/apply-day), les niveaux sont recalculés à partir des "
            "données Waze du jour sélectionné.",
        ),
        (
            "Module incidents (incidents.js)",
            "Les incidents (accident, travaux, météo, autre) sont persistés en base ou en mémoire selon le mode. "
            "Le CRUD complet est exposé en REST ; chaque création, modification ou suppression déclenche un "
            "événement WebSocket pour synchroniser tous les clients connectés. Sur la carte, les incidents "
            "apparaissent avec un marqueur distinct et une fenêtre d'information. Le portail admin permet le "
            "filtrage par statut et la résolution opérationnelle. Ce module illustre la couche « diffusion » "
            "des ITS au sens de Dimitriou.",
        ),
        (
            "Module prédiction (predictionEngine.js, PredictionAnalystPanel)",
            "Pour chaque segment, un historique des niveaux de congestion est maintenu en mémoire. Le moteur "
            "calcule une régression linéaire sur les vingt-quatre derniers points et projette six pas de cinq "
            "minutes (trente minutes d'horizon). Le coefficient R² est renvoyé au client : une valeur proche de 1 "
            "indique une tendance régulière ; une valeur faible signale une prudence accrue dans l'interprétation. "
            "Le panneau analyste agrège les segments les plus congestionnés et les tendances haussières pour "
            "aider l'opérateur du dashboard. L'extension LSTM dans backend/python/ est documentée comme "
            "alignement avec Lv et al. (2015) sans être requise pour valider le PFA.",
        ),
        (
            "Module discover et webcams (discover.js, surveillanceDataset.js, WebcamsPage.jsx)",
            "Après optimisation, le endpoint GET /api/discover/webcams s'appuie exclusivement sur "
            "surveillance_casablanca.json (vingt-cinq points). Les anciennes intégrations ADM, MapCam et "
            "Overpass ont été retirées du chemin critique pour éliminer les timeouts. Le frontend embarque "
            "le dataset via utils/surveillance.js pour un affichage instantané ; la carte Leaflet charge les "
            "tuiles en lazy loading. Les filtres par catégorie (webcam, péage, feu, surveillance) facilitent "
            "la démonstration jury. Les flux vidéo réels, lorsqu'ils existent, s'ouvrent dans un nouvel onglet "
            "plutôt qu'en iframe, pour contourner les restrictions CORS des opérateurs.",
        ),
        (
            "Module chatbot (chatBrain.js, chatAssistant.js, TrafficChatbot.jsx)",
            "L'assistant flottant répond aux questions sur le trafic, les incidents et les fonctionnalités "
            "de l'application. En mode local, des règles et le contexte courant (congestion moyenne, incidents "
            "ouverts) alimentent les réponses sans coût API. Si OPENAI_API_KEY est défini, le backend peut "
            "déléguer à un modèle génératif pour des réponses plus naturelles. Ce module renforce la dimension "
            "« aide à la décision » sans prétendre remplacer un opérateur humain.",
        ),
        (
            "Module logs et alertes (logs.js, alertRoutes)",
            "Les journaux d'activité consultables par l'administrateur tracent connexions, actions sensibles et "
            "erreurs serveur. Les alertes permettent de notifier des seuils de congestion dépassés. Ces briques "
            "préparent une exploitation opérationnelle type centre de contrôle, même si le PFA se concentre sur "
            "la démonstration fonctionnelle plutôt que sur le SLA production.",
        ),
    ]
    for title, body in modules:
        h(doc, title, 2)
        p(doc, body)
        p(
            doc,
            "Tests associés : vérification manuelle via l'interface et, pour l'API, le script smoke "
            "npm run test:api qui contrôle la disponibilité des routes critiques et le code HTTP 200 "
            "sur /api/health.",
        )


def expand_content(doc):
    """Paragraphes additionnels pour épaisseur académique (~pages supplémentaires)."""
    h(doc, "Note sur la reproductibilité", 2)
    p(
        doc,
        "Le dépôt GitHub kalil-cyber/PFA-4eme contient l'intégralité du code, des datasets "
        "(hors fichiers trop volumineux éventuellement en LFS), render.yaml et vercel.json. "
        "Un évaluateur peut cloner, exécuter npm run install:all, lancer backend et frontend, "
        "et reproduire les scénarios des annexes en moins de quinze minutes sur une machine "
        "équipée de Node.js 18+.",
    )
    extensions = [
        (
            "Alignement avec les ODD",
            "L'objectif de développement durable 11 (villes durables) et le 13 (action climatique) "
            "sont indirectement servis par une meilleure fluidité du trafic réduisant les émissions "
            "au ralenti. Tariki ne mesure pas encore le CO₂ évité ; c'est une extension possible "
            "via un module d'impact environnemental.",
        ),
        (
            "Comparaison avec solutions commerciales",
            "Google Maps et Waze offrent une UX mature mais ne fournissent pas de portail "
            "administrateur personnalisable ni d'API ouverte pour Casablanca seule. Tariki se "
            "différencie par la maîtrise du code, l'intégration dataset local et la démonstration "
            "pédagogique des couches ITS.",
        ),
        (
            "Éthique des données",
            "Le dataset utilisé est agrégé et nettoyé ; aucune donnée personnelle nominative "
            "n'est stockée dans les trajectoires exposées. En production, une politique de "
            "confidentialité et une conformité à la loi marocaine sur la protection des données "
            "personnelles seraient requises.",
        ),
        (
            "Travail futur — microservices",
            "Une évolution vers des microservices (auth, trafic, prédiction, discover) permettrait "
            "de scaler indépendamment chaque domaine. Pour un PFA, le monolithe Node reste "
            "justifié par la simplicité de déploiement.",
        ),
    ]
    for title, text in extensions:
        h(doc, title, 3)
        p(doc, text)


def main():
    print("1/2 Génération des schémas PNG…")
    subprocess.run([sys.executable, str(ROOT / "generate_rapport_figures.py")], check=True, cwd=ROOT)

    print("2/2 Composition du document Word…")
    doc = Document()
    setup_document(doc)
    add_footer(doc)

    cover_page(doc)
    front_matter(doc)
    chapter_intro(doc)
    chapter_contexte(doc)
    chapter_etat_art(doc)
    chapter_conception(doc)
    chapter_donnees(doc)
    chapter_realisation(doc)
    chapter_detail_modules(doc)
    expand_content(doc)
    chapter_bilan(doc)
    chapter_biblio_annexes(doc)

    doc.save(str(OUT_DOCX))
    print(f"\n✅ Rapport Word généré :\n   {OUT_DOCX}")
    print("\n📌 Dans Word : clic droit sur la table des matières → « Mettre à jour les champs »")
    print("📌 Export PDF : Fichier → Exporter → PDF (meilleure qualité que l'ancien script fpdf)")


if __name__ == "__main__":
    main()
